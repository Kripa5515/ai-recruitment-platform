import io
import pytest
from docx import Document
from app.ai.extraction.docx_extractor import extract_docx_text
from app.services.exceptions import DOCXExtractionError

def create_test_docx() -> bytes:
    document = Document()
    document.add_paragraph("Kripa Kumar")
    document.add_paragraph("Senior PHP Laravel Developer")
    document.add_paragraph("Python GenAI RAG Developer")
    document.add_paragraph(
        "Skills: PHP, Laravel, Python, PostgreSQL"
    )

    output = io.BytesIO()
    document.save(output)

    return output.getvalue()


def test_extract_docx_text():
    docx_bytes = create_test_docx()

    extracted_text = extract_docx_text(docx_bytes)

    assert "Kripa Kumar" in extracted_text
    assert "Senior PHP Laravel Developer" in extracted_text
    assert "Python GenAI RAG Developer" in extracted_text
    assert "PostgreSQL" in extracted_text

def test_extract_empty_docx_text():
    document = Document()
    output = io.BytesIO()
    document.save(output)
    extracted_text = extract_docx_text(output.getvalue())
    assert extracted_text == ""


def test_extract_invalid_docx():
    invalid_docx = b"this is not a valid docx file"
    with pytest.raises(DOCXExtractionError):
        extract_docx_text(invalid_docx)

def test_extract_docx_table_text():
    document = Document()

    table = document.add_table(rows=2, cols=2)

    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "PHP, Laravel, Python"

    table.cell(1, 0).text = "Database"
    table.cell(1, 1).text = "MySQL, PostgreSQL"

    output = io.BytesIO()
    document.save(output)

    extracted_text = extract_docx_text(output.getvalue())

    assert "Skills" in extracted_text
    assert "PHP, Laravel, Python" in extracted_text
    assert "Database" in extracted_text
    assert "MySQL, PostgreSQL" in extracted_text