import io
from pathlib import Path

import pytest
import pymupdf
from docx import Document

from app.services.exceptions import DuplicateResumeError
from app.services.resume_service import ResumeService


def create_test_pdf(text: str) -> bytes:
    """
    Create a real PDF file in memory for testing.
    """

    document = pymupdf.open()

    page = document.new_page()

    page.insert_text(
        (50, 50),
        text,
    )

    file_content = document.tobytes()

    document.close()

    return file_content


def create_test_docx() -> bytes:
    """
    Create a real DOCX file in memory for testing.
    """

    document = Document()

    document.add_paragraph("Kripa Kumar")
    document.add_paragraph("Senior PHP Laravel Developer")
    document.add_paragraph("Python GenAI RAG Developer")
    document.add_paragraph(
        "Skills: PHP, Laravel, Python, PostgreSQL"
    )

    output = io.BytesIO()

    document.save(output)

    return output.getvalue()


def test_create_resume(db_session):
    service = ResumeService(db_session)

    resume = service.create_resume(
        original_filename="service_resume.pdf",
        file_type="pdf",
        file_size=1024,
        file_hash="1" * 64,
        storage_path="storage/resumes/service_resume.pdf",
    )

    assert resume.id is not None
    assert resume.original_filename == "service_resume.pdf"
    assert resume.file_type == "pdf"
    assert resume.file_size == 1024
    assert resume.file_hash == "1" * 64
    assert resume.source_type == "upload"
    assert resume.extraction_status == "uploaded"


def test_get_resume(db_session):
    service = ResumeService(db_session)

    created = service.create_resume(
        original_filename="resume.pdf",
        file_type="pdf",
        file_size=2000,
        file_hash="2" * 64,
        storage_path="storage/resumes/resume.pdf",
    )

    resume = service.get_resume(created.id)

    assert resume is not None
    assert resume.id == created.id


def test_get_resume_by_hash(db_session):
    service = ResumeService(db_session)

    file_hash = "3" * 64

    created = service.create_resume(
        original_filename="resume.pdf",
        file_type="pdf",
        file_size=3000,
        file_hash=file_hash,
        storage_path="storage/resumes/resume.pdf",
    )

    resume = service.get_resume_by_hash(file_hash)

    assert resume is not None
    assert resume.id == created.id


def test_get_unknown_resume(db_session):
    service = ResumeService(db_session)

    resume = service.get_resume(999999)

    assert resume is None


def test_get_all_resumes(db_session):
    service = ResumeService(db_session)

    service.create_resume(
        original_filename="resume1.pdf",
        file_type="pdf",
        file_size=1000,
        file_hash="4" * 64,
        storage_path="storage/resumes/resume1.pdf",
    )

    service.create_resume(
        original_filename="resume2.docx",
        file_type="docx",
        file_size=2000,
        file_hash="5" * 64,
        storage_path="storage/resumes/resume2.docx",
    )

    resumes = service.get_all_resumes()

    assert len(resumes) >= 2


def test_create_duplicate_resume_hash(db_session):
    service = ResumeService(db_session)

    file_hash = "duplicate" + ("b" * 55)

    service.create_resume(
        original_filename="resume1.pdf",
        file_type="pdf",
        file_size=1000,
        file_hash=file_hash,
        storage_path="storage/resumes/resume1.pdf",
    )

    with pytest.raises(DuplicateResumeError):
        service.create_resume(
            original_filename="resume2.pdf",
            file_type="pdf",
            file_size=2000,
            file_hash=file_hash,
            storage_path="storage/resumes/resume2.pdf",
        )


def test_upload_pdf_resume_creates_resume_and_extracts_text(
    db_session,
    tmp_path,
    monkeypatch,
):
    """
    Test complete PDF upload flow:

    validation
        ↓
    hash calculation
        ↓
    duplicate check
        ↓
    PDF text extraction
        ↓
    file storage
        ↓
    database record
    """

    monkeypatch.setattr(
        "app.core.storage.settings.STORAGE_ROOT",
        str(tmp_path),
    )

    service = ResumeService(db_session)

    file_content = create_test_pdf(
        "Kripa Kumar\n"
        "Senior PHP Laravel Developer\n"
        "Python GenAI RAG Developer"
    )

    resume = service.upload_resume(
        filename="kripa_resume.pdf",
        content_type="application/pdf",
        file_content=file_content,
    )

    assert resume.id is not None

    assert resume.original_filename == "kripa_resume.pdf"

    assert resume.file_type == "pdf"

    assert resume.file_size == len(file_content)

    assert len(resume.file_hash) == 64

    assert resume.storage_path.endswith(".pdf")

    # Verify extracted text
    assert resume.extracted_text is not None

    assert "Kripa Kumar" in resume.extracted_text

    assert "Senior PHP Laravel Developer" in resume.extracted_text

    assert "Python GenAI RAG Developer" in resume.extracted_text

    # Extraction should be completed
    assert resume.extraction_status == "completed"

    # Verify physical file exists
    saved_file = tmp_path / Path(resume.storage_path).name

    assert saved_file.exists()

    # Verify stored file is exactly the original bytes
    assert saved_file.read_bytes() == file_content


def test_upload_duplicate_pdf_resume_is_rejected(
    db_session,
    tmp_path,
    monkeypatch,
):
    """
    Same file content should produce the same SHA-256 hash
    and duplicate upload should be rejected.
    """

    monkeypatch.setattr(
        "app.core.storage.settings.STORAGE_ROOT",
        str(tmp_path),
    )

    service = ResumeService(db_session)

    file_content = create_test_pdf(
        "Duplicate Resume Test"
    )

    # First upload
    first_resume = service.upload_resume(
        filename="kripa_resume.pdf",
        content_type="application/pdf",
        file_content=file_content,
    )

    # Second upload with exactly same content
    with pytest.raises(DuplicateResumeError):
        service.upload_resume(
            filename="another_name.pdf",
            content_type="application/pdf",
            file_content=file_content,
        )

    # Only one physical file should exist
    stored_files = list(tmp_path.iterdir())

    assert len(stored_files) == 1

    # Physical filename should be hash based
    assert (
        stored_files[0].name
        == Path(first_resume.storage_path).name
    )


def test_upload_resume_rejects_unsupported_file(
    db_session,
    tmp_path,
    monkeypatch,
):
    """
    TXT files are not supported.
    """

    monkeypatch.setattr(
        "app.core.storage.settings.STORAGE_ROOT",
        str(tmp_path),
    )

    service = ResumeService(db_session)

    file_content = b"plain text resume"

    with pytest.raises(ValueError):
        service.upload_resume(
            filename="resume.txt",
            content_type="text/plain",
            file_content=file_content,
        )

    # No file should be saved
    assert list(tmp_path.iterdir()) == []

    # No database record should be created
    assert service.get_all_resumes() == []


def test_upload_resume_rejects_invalid_pdf(
    db_session,
    tmp_path,
    monkeypatch,
):
    """
    File has .pdf extension but does not contain
    valid PDF signature.
    """

    monkeypatch.setattr(
        "app.core.storage.settings.STORAGE_ROOT",
        str(tmp_path),
    )

    service = ResumeService(db_session)

    file_content = b"This is not a real PDF file"

    with pytest.raises(ValueError):
        service.upload_resume(
            filename="resume.pdf",
            content_type="application/pdf",
            file_content=file_content,
        )

    # No file should be saved
    assert list(tmp_path.iterdir()) == []

    # No database record should be created
    assert service.get_all_resumes() == []


def test_upload_docx_resume_creates_resume_and_extracts_text(
    db_session,
    tmp_path,
    monkeypatch,
):
    """
    Test complete DOCX upload and text extraction flow.
    """

    monkeypatch.setattr(
        "app.core.storage.settings.STORAGE_ROOT",
        str(tmp_path),
    )

    service = ResumeService(db_session)

    # Create a REAL DOCX file
    file_content = create_test_docx()

    resume = service.upload_resume(
        filename="kripa_resume.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        file_content=file_content,
    )

    assert resume.id is not None

    assert resume.original_filename == "kripa_resume.docx"

    assert resume.file_type == "docx"

    assert resume.file_size == len(file_content)

    assert len(resume.file_hash) == 64

    assert resume.storage_path.endswith(".docx")

    # Verify extracted text
    assert resume.extracted_text is not None

    assert "Kripa Kumar" in resume.extracted_text

    assert "Senior PHP Laravel Developer" in resume.extracted_text

    assert "Python GenAI RAG Developer" in resume.extracted_text

    assert "PostgreSQL" in resume.extracted_text

    # Extraction should be completed
    assert resume.extraction_status == "completed"

    # Verify physical file exists
    saved_file = tmp_path / Path(resume.storage_path).name

    assert saved_file.exists()

    # Verify stored file is exactly the original bytes
    assert saved_file.read_bytes() == file_content


def test_upload_resume_rejects_invalid_docx(
    db_session,
    tmp_path,
    monkeypatch,
):
    """
    File has .docx extension and ZIP signature,
    but is not actually a valid DOCX document.
    """

    monkeypatch.setattr(
        "app.core.storage.settings.STORAGE_ROOT",
        str(tmp_path),
    )

    service = ResumeService(db_session)

    # ZIP signature exists, but this is not a valid DOCX
    file_content = b"PK\x03\x04" + b"invalid docx content"

    with pytest.raises(ValueError):
        service.upload_resume(
            filename="resume.docx",
            content_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            file_content=file_content,
        )

    # Nothing should be stored
    assert list(tmp_path.iterdir()) == []

    # No database record should be created
    assert service.get_all_resumes() == []


def test_get_resume_by_storage_path(db_session):
    service = ResumeService(db_session)

    resume = service.create_resume(
        original_filename="kripa_resume.pdf",
        file_type="pdf",
        file_size=1000,
        file_hash="b" * 64,
        storage_path="storage/resumes/" + "b" * 64 + ".pdf",
    )

    result = service.get_resume_by_storage_path(
        resume.storage_path
    )

    assert result is not None
    assert result.id == resume.id
    assert result.storage_path == resume.storage_path

def test_get_unknown_resume_by_storage_path(db_session):
    service = ResumeService(db_session)

    result = service.get_resume_by_storage_path(
        "storage/resumes/not-found.pdf"
    )

    assert result is None

def test_get_resume_by_source_reference(db_session):
    service = ResumeService(db_session)

    resume = service.create_resume(
        original_filename="kripa_resume.pdf",
        file_type="pdf",
        file_size=1000,
        file_hash="d" * 64,
        storage_path="storage/resumes/" + "d" * 64 + ".pdf",
        source_type="ats",
        source_reference="ATS-67890",
    )

    result = service.get_resume_by_source_reference(
        source_type="ats",
        source_reference="ATS-67890",
    )

    assert result is not None
    assert result.id == resume.id
    assert result.source_reference == "ATS-67890"

def test_get_unknown_resume_by_source_reference(db_session):
    service = ResumeService(db_session)

    result = service.get_resume_by_source_reference(
        source_type="ats",
        source_reference="NOT-FOUND",
    )

    assert result is None

def test_get_or_create_resume_creates_new_resume(
    db_session,
):
    service = ResumeService(db_session)

    resume, created = service.get_or_create_resume(
        original_filename="new_resume.pdf",
        file_type="pdf",
        file_size=1000,
        file_hash="e" * 64,
        storage_path=(
            "storage/resumes/"
            + "e" * 64
            + ".pdf"
        ),
    )

    assert created is True
    assert resume.id is not None
    assert resume.original_filename == "new_resume.pdf"

def test_get_or_create_resume_reuses_existing_resume(
    db_session,
):
    service = ResumeService(db_session)
    first_resume, first_created = (
        service.get_or_create_resume(
            original_filename="resume.pdf",
            file_type="pdf",
            file_size=1000,
            file_hash="f" * 64,
            storage_path=(
                "storage/resumes/"
                + "f" * 64
                + ".pdf"
            ),
        )
    )

    second_resume, second_created = (
        service.get_or_create_resume(
            original_filename="same_resume_different_name.pdf",
            file_type="pdf",
            file_size=1000,
            file_hash="f" * 64,
            storage_path=(
                "storage/resumes/"
                + "f" * 64
                + ".pdf"
            ),
        )
    )

    assert first_created is True
    assert second_created is False

    assert second_resume.id == first_resume.id


def test_ingest_resume_folder_creates_resumes(
    db_session,
    tmp_path,
    monkeypatch,
):
    pdf_file = tmp_path / "kripa.pdf"
    pdf_file.write_bytes(create_test_pdf("Kripa Kumar"))
    docx_file = tmp_path / "john.docx"
    docx_file.write_bytes(create_test_docx())
    
    storage_path = tmp_path / "storage"
    monkeypatch.setattr(
        "app.core.storage.settings.STORAGE_ROOT",
        str(storage_path),
    )

    service = ResumeService(db_session)
    resumes = service.ingest_resume_folder(tmp_path)
    assert len(resumes) == 2
    filenames = {
        resume.original_filename
        for resume in resumes
    }

    assert filenames == {
        "kripa.pdf",
        "john.docx",
    }

def test_ingest_resume_folder_reuses_existing_resumes(
    db_session,
    tmp_path,
    monkeypatch,
):
    pdf_file = tmp_path / "kripa.pdf"
    pdf_file.write_bytes(create_test_pdf("Kripa Kumar"))

    storage_path = tmp_path / "storage"

    monkeypatch.setattr(
        "app.core.storage.settings.STORAGE_ROOT",
        str(storage_path),
    )

    service = ResumeService(db_session)

    first_result = service.ingest_resume_folder(tmp_path)
    second_result = service.ingest_resume_folder(tmp_path)

    assert len(first_result) == 1
    assert len(second_result) == 1

    assert first_result[0].id == second_result[0].id


def test_create_versioned_resume_first_version(db_session):
    from app.data.repositories.candidate_repository import (
        CandidateRepository,
    )

    candidate_repository = CandidateRepository(db_session)
    service = ResumeService(db_session)

    candidate = candidate_repository.create(
        name="Version Candidate",
        email="version1@example.com",
    )

    resume, created = service.create_versioned_resume(
        original_filename="resume_v1.pdf",
        file_type="pdf",
        file_size=1000,
        file_hash="v" * 64,
        storage_path="storage/resumes/" + "v" * 64 + ".pdf",
        candidate_id=candidate.id,
    )

    assert created is True
    assert resume.id is not None
    assert resume.candidate_id == candidate.id
    assert resume.version == 1
    assert resume.is_current is True

def test_create_versioned_resume_creates_next_version(db_session):
    from app.data.repositories.candidate_repository import (
        CandidateRepository,
    )

    candidate_repository = CandidateRepository(db_session)
    service = ResumeService(db_session)

    candidate = candidate_repository.create(
        name="Version Candidate",
        email="version2@example.com",
    )

    resume1, created1 = service.create_versioned_resume(
        original_filename="resume_v1.pdf",
        file_type="pdf",
        file_size=1000,
        file_hash="1" * 64,
        storage_path="storage/resumes/" + "1" * 64 + ".pdf",
        candidate_id=candidate.id,
    )

    resume2, created2 = service.create_versioned_resume(
        original_filename="resume_v2.pdf",
        file_type="pdf",
        file_size=2000,
        file_hash="2" * 64,
        storage_path="storage/resumes/" + "2" * 64 + ".pdf",
        candidate_id=candidate.id,
    )

    assert created1 is True
    assert created2 is True

    assert resume1.version == 1
    assert resume2.version == 2

def test_create_versioned_resume_marks_previous_as_not_current(
    db_session,
):
    from app.data.repositories.candidate_repository import (
        CandidateRepository,
    )

    candidate_repository = CandidateRepository(db_session)
    service = ResumeService(db_session)

    candidate = candidate_repository.create(
        name="Current Test",
        email="current@example.com",
    )

    resume1, _ = service.create_versioned_resume(
        original_filename="resume_v1.pdf",
        file_type="pdf",
        file_size=1000,
        file_hash="a" * 64,
        storage_path="storage/resumes/" + "a" * 64 + ".pdf",
        candidate_id=candidate.id,
    )

    resume2, _ = service.create_versioned_resume(
        original_filename="resume_v2.pdf",
        file_type="pdf",
        file_size=2000,
        file_hash="b" * 64,
        storage_path="storage/resumes/" + "b" * 64 + ".pdf",
        candidate_id=candidate.id,
    )

    db_session.refresh(resume1)
    db_session.refresh(resume2)

    assert resume1.is_current is False
    assert resume2.is_current is True

def test_create_versioned_resume_reuses_duplicate(
    db_session,
):
    from app.data.repositories.candidate_repository import (
        CandidateRepository,
    )

    candidate_repository = CandidateRepository(db_session)
    service = ResumeService(db_session)

    candidate = candidate_repository.create(
        name="Duplicate Version Test",
        email="duplicate-version@example.com",
    )

    file_hash = "d" * 64

    first_resume, first_created = (
        service.create_versioned_resume(
            original_filename="resume_v1.pdf",
            file_type="pdf",
            file_size=1000,
            file_hash=file_hash,
            storage_path="storage/resumes/" + file_hash + ".pdf",
            candidate_id=candidate.id,
        )
    )

    second_resume, second_created = (
        service.create_versioned_resume(
            original_filename="resume_same_content.pdf",
            file_type="pdf",
            file_size=1000,
            file_hash=file_hash,
            storage_path="storage/resumes/" + file_hash + ".pdf",
            candidate_id=candidate.id,
        )
    )

    assert first_created is True
    assert second_created is False

    assert second_resume.id == first_resume.id
    assert second_resume.version == 1

def test_create_versioned_resume_without_candidate(
    db_session,
):
    service = ResumeService(db_session)

    resume, created = service.create_versioned_resume(
        original_filename="unmatched_resume.pdf",
        file_type="pdf",
        file_size=1000,
        file_hash="n" * 64,
        storage_path="storage/resumes/" + "n" * 64 + ".pdf",
        candidate_id=None,
    )

    assert created is True
    assert resume.candidate_id is None
    assert resume.version is None
    assert resume.is_current is True

def test_get_or_create_versioned_resume_from_file_creates_first_version(
    db_session,
):
    from app.data.repositories.candidate_repository import (
        CandidateRepository,
    )

    candidate_repository = CandidateRepository(db_session)

    candidate = candidate_repository.create(
        name="File Version Candidate",
        email="file-version@example.com",
    )

    service = ResumeService(db_session)

    pdf_content = create_test_pdf(
        "Python Laravel PostgreSQL Developer"
    )

    resume, created = (
        service.get_or_create_versioned_resume_from_file(
            filename="candidate_resume.pdf",
            file_content=pdf_content,
            candidate_id=candidate.id,
            content_type="application/pdf",
        )
    )

    assert created is True
    assert resume.candidate_id == candidate.id
    assert resume.version == 1
    assert resume.is_current is True
    assert resume.file_type == "pdf"
    assert resume.extracted_text is not None
    assert "Python" in resume.extracted_text

def test_get_or_create_versioned_resume_from_file_creates_next_version(
    db_session,
):
    from app.data.repositories.candidate_repository import (
        CandidateRepository,
    )

    candidate_repository = CandidateRepository(db_session)

    candidate = candidate_repository.create(
        name="Updated Resume Candidate",
        email="updated-version@example.com",
    )

    service = ResumeService(db_session)

    pdf_v1 = create_test_pdf(
        "Python Developer with Laravel experience"
    )

    resume1, created1 = (
        service.get_or_create_versioned_resume_from_file(
            filename="resume_v1.pdf",
            file_content=pdf_v1,
            candidate_id=candidate.id,
            content_type="application/pdf",
        )
    )

    pdf_v2 = create_test_pdf(
        "Senior Python Developer with Laravel and PostgreSQL"
    )

    resume2, created2 = (
        service.get_or_create_versioned_resume_from_file(
            filename="resume_v2.pdf",
            file_content=pdf_v2,
            candidate_id=candidate.id,
            content_type="application/pdf",
        )
    )

    assert created1 is True
    assert created2 is True

    assert resume1.version == 1
    assert resume2.version == 2

    assert resume1.is_current is False
    assert resume2.is_current is True

    assert resume1.id != resume2.id


def test_get_or_create_versioned_resume_from_file_reuses_duplicate(
    db_session,
):
    from app.data.repositories.candidate_repository import (
        CandidateRepository,
    )

    candidate_repository = CandidateRepository(db_session)

    candidate = candidate_repository.create(
        name="Duplicate File Candidate",
        email="duplicate-file@example.com",
    )

    service = ResumeService(db_session)

    pdf_content = create_test_pdf(
        "Python Developer with PostgreSQL experience"
    )

    resume1, created1 = (
        service.get_or_create_versioned_resume_from_file(
            filename="resume.pdf",
            file_content=pdf_content,
            candidate_id=candidate.id,
            content_type="application/pdf",
        )
    )

    resume2, created2 = (
        service.get_or_create_versioned_resume_from_file(
            filename="resume_copy.pdf",
            file_content=pdf_content,
            candidate_id=candidate.id,
            content_type="application/pdf",
        )
    )

    assert created1 is True
    assert created2 is False

    assert resume2.id == resume1.id
    assert resume2.version == 1
    assert resume2.is_current is True

def test_duplicate_resume_does_not_extract_again(
    db_session,
    monkeypatch,
):
    from app.data.repositories.candidate_repository import (
        CandidateRepository,
    )
    from app.services.resume_service import ResumeService

    candidate_repository = CandidateRepository(db_session)

    candidate = candidate_repository.create(
        name="Incremental Test",
        email="incremental@example.com",
    )

    resume_service = ResumeService(db_session)

    pdf_content = b"%PDF-1.7\nfake resume content"

    extraction_calls = []

    def fake_extract_pdf_text(file_content):
        extraction_calls.append(file_content)
        return "Fake extracted resume text"

    monkeypatch.setattr(
        "app.services.resume_service.extract_pdf_text",
        fake_extract_pdf_text,
    )

    first_resume, first_created = (
        resume_service.get_or_create_versioned_resume_from_file(
            filename="incremental.pdf",
            file_content=pdf_content,
            candidate_id=candidate.id,
        )
    )

    assert first_created is True
    assert len(extraction_calls) == 1

    second_resume, second_created = (
        resume_service.get_or_create_versioned_resume_from_file(
            filename="incremental.pdf",
            file_content=pdf_content,
            candidate_id=candidate.id,
        )
    )

    assert second_created is False
    assert second_resume.id == first_resume.id
    assert second_resume.file_hash == first_resume.file_hash
    assert second_resume.version == first_resume.version

    # Extraction must NOT run again.
    assert len(extraction_calls) == 1