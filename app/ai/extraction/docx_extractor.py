import io
from docx import Document
from app.services.exceptions import DOCXExtractionError

def extract_docx_text(file_content: bytes) -> str:
    """
    Extract text from a DOCX byte stream.

    Args:
        file_content: DOCX file content as bytes.

    Returns:
        Extracted text from paragraphs and tables.

    Raises:
        DOCXExtractionError: If the DOCX cannot be opened or processed.
    """

    try:
        document = Document(io.BytesIO(file_content))

        text_parts = []

        # Extract normal paragraphs
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                text_parts.append(text)

        # Extract table content
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()

                    if text:
                        text_parts.append(text)

        return "\n".join(text_parts).strip()

    except Exception as exc:
        raise DOCXExtractionError(
            "Failed to extract text from DOCX."
        ) from exc